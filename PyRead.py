import pytesseract as pt
from PIL import Image
import tkinter as tk
from tkinter import filedialog
import random
from docx import Document
from docx.shared import Inches
import os


class Utils:
    """Contains the various methods to to convert images to text of different languages and return them in a Word
    doc. """

    def string_to_doc(self, txt, path):
        """Return a .docx file with both the original image and the transcribed text."""
        result = Document()
        result.add_heading('Results of OCR conversion', 0)
        for i in range(len(path)):
            result.add_picture(path[i], width=Inches(5.0))
            result.add_paragraph(txt[i])
        file_location_and_name = filedialog.asksaveasfilename(initialdir="/",
                                                              title="Select file location and input file name.",
                                                              filetypes=([("DOCs", "*.docx .doc")]))
        chosen_name = os.path.basename(file_location_and_name)
        chosen_path = os.path.dirname(file_location_and_name)
        result.save(os.path.join(chosen_path, chosen_name))

    def img_to_text(self, images, file_path, language):
        """Return an array of strings of all OCR-translated files."""
        resultant_text = []
        for img in images:
            ocr_readable_eng_img = Image.open(img)
            resultant_text.append(pt.image_to_string(ocr_readable_eng_img, lang=language))
        Utils.string_to_doc(self, resultant_text, file_path)
        return resultant_text


if __name__ == '__main__':

    obj = Utils()
    while True:
        selection = input('Which language would you like to convert your image to text? 1. English 2. Spanish 3. '
                          'French 4. German 5. Japanese 6. Simplified Chinese 7. Exit\n')
        root = tk.Tk()
        root.withdraw()
        image_selection, selection_path = tk.filedialog.askopenfilenames(title='Choose your images.',
                                                                         filetypes=[("PNG", "*.png"),
                                                                                    ("JPEGs", "*.jpeg jpg"),
                                                                                    ("GIF", "*.gif"), (("BMP", "*.bmp"),
                                                                                                       ("tiff",
                                                                                                        "*.tif tiff"))])
        if selection not in "1234567":
            print('Invalid choice.')
            continue
        elif selection == '1':
            obj.img_to_text(image_selection, selection_path, 'eng')
        elif selection == '2':
            obj.img_to_text(image_selection, selection_path, 'spa')
        elif selection == '3':
            obj.img_to_french(image_selection, selection_path, 'fra')
        elif selection == '4':
            obj.img_to_german(image_selection, selection_path, 'deu')
        elif selection == '5':
            obj.img_to_japanese(image_selection, selection_path, 'jpn')
        elif selection == '6':
            obj.img_to_chinese(image_selection, selection_path, 'chi_sim')
        else:
            goodbye = ['¡Adiós!', 'Goodbye!', 'じゃね。', 'Tschüss!', 'Au revoir!', '再见。']
            print(random.choice(goodbye))
            break
