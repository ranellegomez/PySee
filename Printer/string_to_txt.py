import PyRead
import io
import collections
from docx import newdocument


def string_to_doc(txt, images=None):
    """Takes a string from one of the Scanner methods and returns a doc with
    both the original image and the transcribed text."""
    result = newdocument()
    result.add_heading('Results of OCR scan.', 0)
    if not isinstance(images, collections.iterable):
        result.add_picture(img, width=Inches(1.25))
    else:
        for img in images:
            result.add_picture(img)
            result.add_paragraph(txt)
    result.save('translation.docx')


